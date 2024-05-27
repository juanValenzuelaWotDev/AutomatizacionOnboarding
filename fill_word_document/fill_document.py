# Import libraries
from docx import Document
from docx2pdf import convert
import json
import os
import sys
# Take the main path of AutomatizacionOnboarding and use it as normal
sys.path.insert(1, "/".join(os.path.realpath(__file__).split("/")[0:-2]))

import bot_paths

# turn numbers to words for contract fill
from num2words import num2words

# Global variables
# ----------------------------------------------------------------------------------------------------
# Ejemplo de datos de WOT
data_to_replace_empleado = {
    'nombre_empleado' : 'NOMBRE_EMPLEADO',
    'anio': 'FECHA_ANIO',
    'mes': 'FECHA_MES',
    'dia': 'FECHA_DIA',
    'puesto': 'PUESTO_EMPLEADO',
    'cliente': 'CLIENTE_PROYECTO',
    'dia_inicio': 'DIA_INICIO',
    'mes_inicio': 'MES_INICIO',
    'anio_inicio': 'ANIO_INICIO',
    'duracion_puesto': 'DURACION_PUESTO',
    'pago_texto_hora': 'PAGO_TEXTO_HORA',
    'pago_hora': 'PAGO_NUMERO_HORA',
    'pago_texto_mes': 'PAGO_TEXTO_MES',
    'pago_mes': 'PAGO_NUMERO_MES',
    'pago_texto_mes_quetzales': 'PAGO_QUETZALES_TEXTO_MES',
    'pago_mes_quetzales': 'PAGO_QUETZALES_NUMERO_MES',
    'modalidad': 'MODALIDAD_TRABAJO',
}



# Functions that turns integrers into words
# https://es.stackoverflow.com/questions/584184/convertir-n%C3%BAmeros-a-palabras-con-python    
def simplify_num(num: str):
    """
    This function divides the number in chunks of 3 digits
    These chunks are much easier to process and assign the correct word
    """
    cui_number = num[:]
    
    # first erase the spaces between numbers
    if " " in cui_number:
        return cui_number.split(" ")

    # next split the numbers into three standard chunks
    return [cui_number[0:4], cui_number[4:9], cui_number[9:14]]

def num_to_words(number:str):
    num_list = simplify_num(number)
    num_buffer = []
    word_buffer = []
    for item in num_list:
        if item[0] == '0':
            num_buffer.append(int(item[0]))
            num_buffer.append(int(item[1:]))
            word_buffer.append(num2words(int(item[0]), lang='es'))
            word_buffer.append(num2words(int(item[1:]),lang='es'))
        else:
            num_buffer.append(int(item))
            word_buffer.append(num2words(int(item),lang='es'))    
    # print(num_buffer)
    # print(word_buffer)    
    # Return a single string containing all the numbers in word form
    return ', '.join(word_buffer)


# This function replaces one word
def replace_one_word(file_path, search_word, replace_word):
    """
    This function replaces one word in a document
    file path is the path where it looks for the template
    search word is the word that this looks for
    replace word is the word that this puts in place of the search word
    """

    # # Open an existing document
    print("opening document...")
    doc = Document(f'.../downloads/{file_path}')

    for paragraph in doc.paragraphs:
        if search_word in paragraph.text:
            # this stores the format of the text
            inline = paragraph.runs
            for i in range(len(inline)):
                if search_word in inline[i].text:
                    text= inline[i].text.replace(search_word, replace_word)
                    inline[i].text = text


    # # Save the modified document
    print("saving document...")
    doc.save('modified_document.docx')
    print("saved!")
    # in here just delete the input file so it doesn't duplicate the next run

# works for multiple words in one document
def fill_document(file_path, search_data:dict, replace_data:dict, type:str="wot"):
    """
    This function replaces multiple keywords in a document 
    file path is the path where it looks for the template
    search word is the word that this looks for
    replace word is the word that this puts in place of the search word
    """
    # Download contract
    if type.lower()=="wot":
        # download_contract()
        print("descargando archivo de wot")
    elif type.lower()=="bot":
        print("descargando archivo de bot")
    else:
        print("tipo de contrato no valido")
    # # Open an existing document
    print("opening document...")
    doc = Document(f'{bot_paths.download_path}/{file_path}')
    for k in search_data:
        search_word = search_data[k]
        replace_word = replace_data[k]
        print(f"Key: {k}")
        print(f"Look for: {search_word}")
        print(f"Replace with: {replace_word}")
        for paragraph in doc.paragraphs:
            if search_word in paragraph.text:
                # this stores the format of the text
                inline = paragraph.runs
                for i in range(len(inline)):
                    if search_word in inline[i].text:
                        text= inline[i].text.replace(search_word, replace_word)
                        inline[i].text = text
                        print('replaced')


    # # Save the modified document
    print("saving document...")
    temporary_name = r"{}\modified_document.docx".format(bot_paths.modified_document_path)
    doc.save(temporary_name)
    # Remove old extension and replace with .pdf
    new_name = file_path.split(".")[0]
    new_name = new_name.replace(" [Template]","")
    # Replace all accented characters (tildes)
    if ("á" in new_name):
        new_name = new_name.replace("á","a")
    elif ("é" in new_name):
        new_name = new_name.replace("é","e")
    elif ("í" in new_name):
        new_name = new_name.replace("í","i")
    elif ("ó" in new_name):
        new_name = new_name.replace("ó","o")
    elif ("ú" in new_name):
        new_name = new_name.replace("ú","u")
    # Add more letters as needed

    convert(temporary_name, r'{}\{}.pdf'.format(bot_paths.modified_document_path, new_name))
    # Delete the temporary docx file
    if (os.path.exists(temporary_name)):
        os.remove(temporary_name)
        # clean_dir(picture_path) #clean the download directory to avoid errors
    print("saved!")

    return f"{new_name}.pdf"


